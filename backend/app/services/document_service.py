"""
Document Service - Orchestrates Document Processing Pipeline
Upload → S3 → OCR (Textract / docx parser) → Entity Extraction (Comprehend) → Classification (Bedrock) → RAG Store
"""
import io
import logging
import json
from decimal import Decimal
from typing import Dict, List, Optional, Any
from app.services.s3_service import s3_service
from app.services.textract_service import textract_service
from app.services.comprehend_service import comprehend_service
from app.services.bedrock_service import bedrock_service
from app.services.dynamodb_service import db
from app.utils.helpers import generate_id, file_hash, get_document_category, now_iso

logger = logging.getLogger(__name__)


def _extract_docx_text(file_content: bytes) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(file_content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed; cannot extract text from .docx files")
        return ""
    except Exception as e:
        logger.error(f"docx text extraction failed: {e}")
        return ""


def _sanitize_for_dynamo(obj: Any) -> Any:
    """Recursively convert float → Decimal for DynamoDB compatibility."""
    if isinstance(obj, float):
        return Decimal(str(obj))
    elif isinstance(obj, dict):
        return {k: _sanitize_for_dynamo(v) for k, v in obj.items()}
    elif isinstance(obj, (list, tuple)):
        return [_sanitize_for_dynamo(i) for i in obj]
    return obj


class DocumentService:
    """Full document processing pipeline"""
    
    def process_document(self, user_id: str, file_content: bytes, filename: str,
                         content_type: str, document_type: str = None) -> Dict:
        """Complete pipeline: Upload → OCR → Extract → Classify → Store"""
        
        document_id = generate_id()
        content_hash = file_hash(file_content)
        
        # Step 1: Check for duplicates
        existing_docs = db.get_user_documents(user_id)
        for doc in existing_docs:
            if doc.get("content_hash") == content_hash:
                return {
                    "status": "duplicate",
                    "message": "This document has already been uploaded",
                    "existing_document_id": doc["document_id"]
                }
        
        # Step 2: Upload to S3
        ext = filename.rsplit(".", 1)[-1] if "." in filename else "pdf"
        s3_key = f"documents/{user_id}/{document_id}.{ext}"
        s3_service.upload_file(file_content, s3_key, content_type)
        
        # Step 3: OCR / text extraction
        # .docx files → python-docx; images/PDFs → AWS Textract
        ext_lower = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
        is_docx = ext_lower in ("docx", "doc") or "wordprocessingml" in (content_type or "")
        try:
            if is_docx:
                ocr_text = _extract_docx_text(file_content)
                ocr_confidence = 100 if ocr_text else 0
            else:
                ocr_result = textract_service.extract_text(file_content)
                ocr_text = ocr_result.get("full_text", "")
                ocr_confidence = ocr_result.get("confidence", 0)
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            ocr_text = ""
            ocr_confidence = 0
        
        # Step 4: Entity Extraction with Comprehend
        extracted_entities = {}
        if ocr_text:
            try:
                entity_result = comprehend_service.extract_entities(ocr_text)
                extracted_entities = entity_result.get("entities", {})
            except Exception as e:
                logger.error(f"Entity extraction failed: {e}")
        
        # Step 5: Classification with Bedrock
        classified_type = document_type or "other"
        ai_generated_name = filename
        extracted_data = {}
        
        if ocr_text:
            try:
                classification = bedrock_service.classify_document(ocr_text)
                if not document_type:
                    classified_type = classification.get("document_type", "other")
                extracted_data = classification.get("extracted_data", {})
                ai_generated_name = classification.get("ai_generated_name", filename)
            except Exception as e:
                logger.error(f"Classification failed: {e}")
        
        # Step 6: Store metadata in DynamoDB
        # DynamoDB does not support Python floats — use Decimal or str for all numeric fields
        doc_data = {
            "user_id": user_id,
            "document_id": document_id,
            "original_filename": filename,
            "ai_generated_name": ai_generated_name,
            "document_type": classified_type,
            "category": get_document_category(classified_type),
            "s3_key": s3_key,
            "content_type": content_type,
            "file_size": len(file_content),
            "content_hash": content_hash,
            "status": "processed",
            # Store full OCR text for RAG (DynamoDB item limit is 400KB - truncate at 50k chars)
            "ocr_text": ocr_text[:50000] if ocr_text else "",
            "ocr_confidence": str(ocr_confidence),  # store as string to avoid float
            # Sanitize nested dicts/floats for DynamoDB
            "extracted_data": _sanitize_for_dynamo(extracted_data),
            "extracted_entities": _sanitize_for_dynamo(extracted_entities),
            "is_verified": False,
            "upload_date": now_iso(),
            "processed_date": now_iso(),
        }
        
        db.save_document(doc_data)
        
        # Generate presigned URL for viewing
        try:
            view_url = s3_service.get_presigned_url(s3_key)
        except Exception:
            view_url = ""
        
        return {
            "status": "processed",
            "document_id": document_id,
            "document_type": classified_type,
            "ai_generated_name": ai_generated_name,
            "extracted_data": extracted_data,
            "ocr_confidence": ocr_confidence,
            "ocr_text_preview": ocr_text[:200] if ocr_text else "",
            "view_url": view_url,
        }
    
    def get_document(self, user_id: str, document_id: str) -> Optional[Dict]:
        """Get document with presigned URL"""
        doc = db.get_document(user_id, document_id)
        if doc and doc.get("s3_key"):
            doc["view_url"] = s3_service.get_presigned_url(doc["s3_key"])
        return doc
    
    def get_user_documents(self, user_id: str) -> list:
        """Get all documents for user with presigned URLs"""
        docs = db.get_user_documents(user_id)
        for doc in docs:
            if doc.get("s3_key"):
                doc["view_url"] = s3_service.get_presigned_url(doc["s3_key"])
        return docs
    
    def delete_document(self, user_id: str, document_id: str) -> bool:
        """Delete document from S3 and DynamoDB"""
        doc = db.get_document(user_id, document_id)
        if not doc:
            return False
        
        # Delete from S3
        if doc.get("s3_key"):
            s3_service.delete_file(doc["s3_key"])
        
        # Delete from DynamoDB
        return db.delete_document(user_id, document_id)
    
    def get_user_document_context(self, user_id: str) -> str:
        """
        Build a RAG context string from all user's uploaded documents.
        Used by the AI chat and form-fill agents to access the user's
        personal document data (name, Aadhaar number, DOB, address, etc.)
        """
        try:
            docs = db.get_user_documents(user_id)
        except Exception:
            return ""

        if not docs:
            return ""

        context_parts = ["[User's Uploaded Documents — use this information for form filling and eligibility matching:]"]

        for doc in docs:
            doc_type = doc.get("document_type", "unknown")
            ai_name = doc.get("ai_generated_name", doc.get("original_filename", "unknown"))
            extracted = doc.get("extracted_data", {})
            ocr_text = doc.get("ocr_text", "")

            section = [f"\n## Document: {ai_name} (Type: {doc_type})"]

            # Include structured extracted fields first (highest quality)
            if extracted:
                fields = []
                for key, val in extracted.items():
                    if key == "other_fields" and isinstance(val, dict):
                        for k2, v2 in val.items():
                            if v2:
                                fields.append(f"  - {k2}: {v2}")
                    elif val:
                        fields.append(f"  - {key}: {val}")
                if fields:
                    section.append("  Extracted Fields:")
                    section.extend(fields)

            # Include raw OCR text (trimmed) for any additional info
            if ocr_text:
                # Limit OCR snippet to 1000 chars per document to keep context manageable
                section.append(f"  OCR Text (first 1000 chars): {ocr_text[:1000]}")

            context_parts.append("\n".join(section))

        return "\n".join(context_parts)

    def get_document_map_for_form(self, user_id: str) -> Dict:
        """
        Returns a flat dict of key user data fields extracted from all documents.
        Used by form-fill agent to auto-populate form fields.
        """
        try:
            docs = db.get_user_documents(user_id)
        except Exception:
            return {}

        merged = {}
        # Process in order: identity docs first, then others
        priority_types = ["aadhaar", "pan", "voter_id", "passport", "driving_license"]
        sorted_docs = sorted(
            docs,
            key=lambda d: (priority_types.index(d.get("document_type", "")) 
                           if d.get("document_type", "") in priority_types else 99)
        )

        for doc in sorted_docs:
            extracted = doc.get("extracted_data", {})
            if isinstance(extracted, dict):
                for key, val in extracted.items():
                    if key == "other_fields" and isinstance(val, dict):
                        for k2, v2 in val.items():
                            if v2 and k2 not in merged:
                                merged[k2] = v2
                    elif val and key not in merged:
                        merged[key] = val

        return merged

    def check_required_documents(self, user_id: str, required_docs: list) -> Dict:
        """Check which required documents user has"""
        user_docs = db.get_user_documents(user_id)
        user_doc_types = {doc["document_type"] for doc in user_docs}
        
        available = []
        missing = []
        
        for req_doc in required_docs:
            doc_type = req_doc if isinstance(req_doc, str) else req_doc.get("document_type", "")
            if doc_type in user_doc_types:
                available.append(doc_type)
            else:
                missing.append(doc_type)
        
        return {
            "available": available,
            "missing": missing,
            "all_available": len(missing) == 0,
            "total_required": len(required_docs),
            "total_available": len(available)
        }


# Singleton
document_service = DocumentService()
